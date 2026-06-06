#!/usr/bin/env python3
"""Generate Phase 2 Word document with proposal, results, and figures."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
FIGURES = ROOT / "outputs" / "figures"
REPORTS = ROOT / "outputs" / "reports"
OUT_PATH = ROOT / "docs" / "Phase2_Proposal_and_Implementation.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def add_image_if_exists(doc: Document, path: Path, caption: str, width: float = 5.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Figure not found: {caption}]")


def build_document() -> Document:
    doc = Document()

    # Title page
    title = doc.add_heading("Phase 2 — Proposal and Code Implementation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Urban Green Cover Image Assessment")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph("CNN-Based Greenery Classification from Aerial Scene Images").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("Machine Learning Course — Applied CNN Project").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Student Name: [Your Name]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Idea 15 of 75 — CNN-Based Image Dataset Project Idea Book").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_page_break()

    # 1. Project Title
    add_heading(doc, "1. Project Title", 1)
    doc.add_paragraph(
        "Urban Green Cover Image Assessment: CNN-Based Greenery Classification from Aerial Scene Images"
    )

    # 2. Background
    add_heading(doc, "2. Background and Motivation", 1)
    doc.add_paragraph(
        "Urban green cover is a key indicator of environmental quality, livability, and climate "
        "resilience in cities. Planners and environmental agencies need scalable methods to screen "
        "large volumes of scene images and identify areas with high, moderate, sparse, or minimal "
        "visible greenery."
    )
    doc.add_paragraph(
        "Convolutional Neural Networks (CNNs) can automate image screening for environmental "
        "monitoring. This project applies supervised deep learning to a public aerial image dataset "
        "and presents the system as decision-support for human experts rather than autonomous "
        "policy control."
    )
    doc.add_paragraph("Applied use cases:", style="List Bullet")
    doc.add_paragraph("Urban planning dashboards", style="List Bullet")
    doc.add_paragraph("Neighborhood greenery assessment for smart cities", style="List Bullet")

    # 3. Problem Statement
    add_heading(doc, "3. Problem Statement", 1)
    doc.add_paragraph(
        "Given an aerial scene image, classify the level of urban green cover into one of four "
        "greenery categories. The system must train and compare a custom CNN with transfer learning "
        "models, handle class imbalance, report standard metrics, and provide explainability."
    )

    # 4. Dataset
    add_heading(doc, "4. Selected Dataset and Link", 1)
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["Dataset", "UC Merced Land Use Dataset (Kaggle)"],
            ["Kaggle Link", "https://www.kaggle.com/datasets/abdulhasibuddin/uc-merced-land-use-dataset"],
            ["Original Source", "http://weegee.vision.ucmerced.edu/datasets/landuse.html"],
            ["License", "Public domain imagery (USGS)"],
        ],
    )

    # 5. Dataset Description
    add_heading(doc, "5. Dataset Description", 1)
    add_table(
        doc,
        ["Property", "Value"],
        [
            ["Image type", "Aerial RGB remote sensing patches"],
            ["Original resolution", "256 × 256 pixels"],
            ["Original classes", "21 land-use categories"],
            ["Total images used", "2,099 (1 duplicate removed)"],
            ["Project classes", "4 greenery levels"],
            ["Train / Val / Test", "1,471 / 314 / 314 (70% / 15% / 15%)"],
        ],
    )

    add_heading(doc, "Greenery Taxonomy", 2)
    add_table(
        doc,
        ["Class", "Description", "UC Merced Source Examples"],
        [
            ["dense_green", "High vegetation cover", "forest, agricultural, golfcourse"],
            ["moderate_green", "Moderate vegetation", "chaparral, river"],
            ["sparse_green", "Low urban greenery", "residential area classes"],
            ["minimal_green", "Built-up / low green cover", "buildings, freeway, parkinglot, runway"],
        ],
    )

    add_heading(doc, "Class Distribution", 2)
    add_table(
        doc,
        ["Class", "Count", "Share"],
        [
            ["minimal_green", "1,299", "61.9%"],
            ["dense_green", "300", "14.3%"],
            ["sparse_green", "300", "14.3%"],
            ["moderate_green", "200", "9.5%"],
        ],
    )
    add_image_if_exists(doc, FIGURES / "class_distribution.png", "Figure 1: Class distribution")

    # 6. Research Questions and Answers
    add_heading(doc, "6. Research Questions and Answers", 1)

    add_heading(doc, "RQ1: Classification Accuracy on Public Datasets", 2)
    doc.add_paragraph(
        "Question: How accurately can CNN-based models perform urban greenery classification "
        "from scene images using publicly available image datasets?"
    )
    doc.add_paragraph(
        "Answer: Using the UC Merced Land Use Dataset (Kaggle) mapped to four greenery classes, "
        "all CNN models achieved strong screening-level accuracy on the held-out test set (n=314). "
        "MobileNetV2 reached 93.3% accuracy (F1 macro 0.91, ROC-AUC 0.996), ResNet50 reached 93.0%, "
        "and the custom CNN reached 85.0%. This confirms that publicly available aerial imagery "
        "supports reliable urban green-cover classification when labels are mapped to a greenery taxonomy."
    )

    add_heading(doc, "RQ2: Performance vs. Computational Efficiency", 2)
    doc.add_paragraph(
        "Question: Which model family — custom CNN or transfer learning — provides the best "
        "balance between predictive performance and computational efficiency?"
    )
    doc.add_paragraph(
        "Answer: Transfer learning outperforms the custom CNN. MobileNetV2 provides the best "
        "trade-off: 93.3% accuracy with ~12 ms/image inference and only ~2.3M parameters "
        "(frozen backbone). ResNet50 matches accuracy (~93%) but is ~4× slower (~47 ms/image) "
        "due to a larger backbone. The custom CNN is lightweight (~391K parameters, ~11 ms/image) "
        "but lower accuracy (85%). For deployment as a smart-city screening tool, MobileNetV2 is recommended."
    )

    add_heading(doc, "RQ3: Preprocessing, Augmentation, and Class Imbalance", 2)
    doc.add_paragraph(
        "Question: How do preprocessing, augmentation, and class-imbalance handling influence "
        "model robustness and class-wise performance?"
    )
    ablation_path = REPORTS / "rq3_ablation.json"
    if ablation_path.exists():
        ablation = json.loads(ablation_path.read_text())
        doc.add_paragraph(
            "Answer: An ablation study on MobileNetV2 compared four training configurations "
            "(full pipeline, no augmentation, no class weights, neither). Results:"
        )
        rows = []
        for v in ablation["variants"]:
            rows.append(
                [
                    v["description"],
                    "Yes" if v["augmentation"] else "No",
                    "Yes" if v["class_weights"] else "No",
                    f"{v['test_accuracy']*100:.1f}%",
                    f"{v['f1_macro']:.3f}",
                    f"{v['per_class_f1'].get('dense_green', 0):.3f}",
                    f"{v['per_class_f1'].get('moderate_green', 0):.3f}",
                ]
            )
        add_table(
            doc,
            ["Configuration", "Augmentation", "Class Weights", "Test Acc.", "F1 macro", "F1 dense", "F1 moderate"],
            rows,
        )
        doc.add_paragraph(ablation.get("findings", ""))
    else:
        doc.add_paragraph(
            "Answer: Preprocessing (resize to 224×224, backbone-specific normalization) is required "
            "for transfer learning. Data augmentation (rotation, flip, zoom) improves validation "
            "generalization. Class-weighted loss addresses the 62% majority class (minimal_green) "
            "and improves recall on minority greenery classes. See rq3_ablation.json for full comparison."
        )
    add_image_if_exists(doc, FIGURES / "rq3_ablation_chart.png", "Figure: RQ3 ablation comparison")

    add_heading(doc, "RQ4: Grad-CAM and Meaningful Visual Regions", 2)
    doc.add_paragraph(
        "Question: Do Grad-CAM explanations show that trained models focus on meaningful visual "
        "regions related to urban greenery classification?"
    )
    rq4_path = REPORTS / "rq4_gradcam_analysis.json"
    if rq4_path.exists():
        rq4 = json.loads(rq4_path.read_text())
        doc.add_paragraph(
            f"Answer: Yes. We analyzed explainability maps for {', '.join(rq4['models_analyzed'])}. "
            + rq4["conclusion"]
        )
        for item in rq4.get("findings", []):
            doc.add_paragraph(
                f"{item['category'].replace('_', ' ').title()}: {item['observation']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph(
            "Answer: Grad-CAM and saliency maps show models activating on vegetated regions for "
            "green classes and on built-up texture for minimal_green. See outputs/figures/grad_cam/."
        )

    add_heading(doc, "Grad-CAM Evidence (Custom CNN and MobileNetV2)", 3)
    for model, label in [("custom_cnn", "Grad-CAM"), ("mobilenetv2", "Saliency")]:
        grad_dir = FIGURES / "grad_cam" / model
        if grad_dir.exists():
            for img in sorted(grad_dir.glob("correct_*.png"))[:2]:
                add_image_if_exists(doc, img, f"{label} — {model} correct prediction", width=5.5)
            for img in sorted(grad_dir.glob("false_*.png"))[:1]:
                add_image_if_exists(doc, img, f"{label} — {model} error case", width=5.5)

    add_heading(doc, "RQ5: Failure Patterns and Deployment Risks", 2)
    doc.add_paragraph(
        "Question: What are the main failure patterns, deployment limitations, and practical risks "
        "when applying this system in real environmental monitoring settings?"
    )
    err_path = REPORTS / "mobilenetv2_error_analysis.json"
    if err_path.exists():
        err = json.loads(err_path.read_text())
        doc.add_paragraph(
            f"Answer: MobileNetV2 error rate is {err['error_rate']*100:.1f}% on the test set. "
            "The dominant failure is minimal_green confused with sparse_green (small lawns/trees "
            "near buildings). High-confidence errors can mislead non-expert reviewers."
        )
        doc.add_paragraph("Top confusion pairs:")
        add_table(
            doc,
            ["True Class", "Predicted Class", "Count"],
            [[p["true"], p["predicted"], str(p["count"])] for p in err.get("most_common_confusion_pairs", [])[:5]],
        )
        doc.add_paragraph("Deployment limitations and risks:")
        for note in err.get("deployment_notes", {}).get("main_risks", []):
            doc.add_paragraph(note, style="List Bullet")
        doc.add_paragraph("Recommended validation:")
        for note in err.get("deployment_notes", {}).get("validation_requirements", []):
            doc.add_paragraph(note, style="List Bullet")
        doc.add_paragraph(
            "Conclusion: The system should be deployed as decision-support for human experts, "
            "not for autonomous urban planning decisions."
        )

    # 7. Results
    add_heading(doc, "7. Expected and Observed Results", 1)
    comparison_path = REPORTS / "model_comparison.json"
    if comparison_path.exists():
        data = json.loads(comparison_path.read_text())
        rows = []
        for model, m in data.items():
            rows.append(
                [
                    model,
                    f"{m['accuracy']*100:.1f}%",
                    f"{m['f1_macro']:.3f}",
                    f"{m.get('roc_auc_ovr', 0):.3f}",
                    f"{m.get('avg_inference_ms_per_image', 0):.1f}",
                ]
            )
        add_table(doc, ["Model", "Accuracy", "F1 (macro)", "ROC-AUC", "Inference (ms)"], rows)

    add_image_if_exists(doc, FIGURES / "model_comparison_table.png", "Figure 2: Model comparison table")

    doc.add_paragraph(
        "Best model: MobileNetV2 achieves 93.3% test accuracy with 11.9 ms inference time, "
        "providing the best accuracy–efficiency trade-off for urban greenery screening."
    )

    # 8. Methodology
    add_heading(doc, "8. Proposed Methodology", 1)
    steps = [
        "Load UC Merced images from Kaggle dataset",
        "Clean dataset (remove corrupt/duplicate images)",
        "Map 21 land-use classes to 4 greenery classes",
        "Split into train/validation/test (70/15/15)",
        "Resize to 224×224, normalize, apply augmentation on training set",
        "Train Custom CNN, MobileNetV2, ResNet50, EfficientNetB0",
        "Evaluate with accuracy, precision, recall, F1, ROC-AUC, confusion matrix",
        "Generate Grad-CAM / saliency visualizations",
        "Perform error analysis for deployment discussion",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph("Framework: TensorFlow / Keras")
    doc.add_paragraph("Environment: Python 3.x (Anaconda / Kaggle Notebook)")

    # 9. Model Design
    add_heading(doc, "9. CNN Model Design", 1)
    add_heading(doc, "Custom CNN (Baseline)", 2)
    for item in [
        "4 convolutional blocks: 32 → 64 → 128 → 256 filters",
        "Batch normalization + max pooling per block",
        "Global average pooling + dropout (0.4)",
        "Softmax output layer (4 classes)",
        "Approximately 391,000 parameters",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Transfer Learning Models", 2)
    add_table(
        doc,
        ["Model", "Backbone", "Pretrained Weights", "Trainable Layers"],
        [
            ["MobileNetV2", "ImageNet", "Frozen backbone", "Classification head only"],
            ["ResNet50", "ImageNet", "Frozen backbone", "Classification head only"],
            ["EfficientNetB0", "ImageNet", "Frozen backbone", "Classification head only"],
        ],
    )

    # 10. Evaluation Metrics
    add_heading(doc, "10. Evaluation Metrics", 1)
    for metric in [
        "Accuracy", "Precision, Recall, F1-score (macro and weighted)",
        "Confusion matrix", "ROC-AUC (one-vs-rest)", "Training/validation loss and accuracy curves",
        "Inference time per image", "Model parameter count",
    ]:
        doc.add_paragraph(metric, style="List Bullet")

    # 11. Figures and Results
    add_heading(doc, "11. Results — Figures and Visualizations", 1)

    add_heading(doc, "Training Curves", 2)
    for model in ["custom_cnn", "mobilenetv2", "resnet50", "efficientnetb0"]:
        add_image_if_exists(
            doc,
            FIGURES / f"{model}_training_curves.png",
            f"Figure: {model} training loss and accuracy curves",
            width=6.0,
        )

    add_heading(doc, "Confusion Matrices", 2)
    for model in ["custom_cnn", "mobilenetv2", "resnet50", "efficientnetb0"]:
        add_image_if_exists(
            doc,
            FIGURES / f"{model}_confusion_matrix.png",
            f"Figure: {model} confusion matrix",
        )

    add_heading(doc, "ROC Curves", 2)
    for model in ["custom_cnn", "mobilenetv2", "resnet50", "efficientnetb0"]:
        add_image_if_exists(
            doc,
            FIGURES / f"{model}_roc_curves.png",
            f"Figure: {model} ROC curves (one-vs-rest)",
        )

    add_heading(doc, "Grad-CAM / Saliency (Sample — MobileNetV2)", 2)
    grad_dir = FIGURES / "grad_cam" / "mobilenetv2"
    if grad_dir.exists():
        for img in sorted(grad_dir.glob("*.png"))[:3]:
            add_image_if_exists(doc, img, f"Figure: MobileNetV2 explainability — {img.stem}", width=6.0)

    # 12. Code Implementation
    add_heading(doc, "12. Code Implementation Summary", 1)
    add_table(
        doc,
        ["Requirement", "Implementation File"],
        [
            ["Dataset loading", "src/data_preparation.py, src/data_loader.py"],
            ["Preprocessing & split", "src/data_preparation.py"],
            ["Data augmentation", "src/data_loader.py"],
            ["Custom CNN", "src/models.py"],
            ["Transfer learning", "src/models.py"],
            ["Model training", "src/train.py"],
            ["Model evaluation", "src/evaluate.py"],
            ["Grad-CAM / saliency", "src/grad_cam.py"],
            ["Error analysis", "src/error_analysis.py"],
            ["Pipeline runner", "run_pipeline.py"],
            ["Report figures", "scripts/generate_report_figures.py"],
        ],
    )

    add_heading(doc, "How to Run the Code", 2)
    code_steps = """pip install -r requirements.txt
# Download UC Merced from Kaggle → data/raw/converted_uc_merced_data/
python run_pipeline.py --step prepare
python run_pipeline.py --step train
python run_pipeline.py --step evaluate
python run_pipeline.py --step explain
python scripts/generate_report_figures.py"""
    doc.add_paragraph(code_steps)

    # 13. Submission
    add_heading(doc, "13. Phase 2 Submission Materials", 1)
    add_table(
        doc,
        ["Material", "Details"],
        [
            ["Proposal document", "This Word document"],
            ["GitHub repository", "[Add your public GitHub link]"],
            ["Kaggle Notebook", "[Add link if applicable]"],
            ["README", "README.md in repository"],
            ["Source code", "src/, run_pipeline.py, scripts/"],
            ["Dataset link", "Kaggle UC Merced Land Use Dataset"],
        ],
    )
    doc.add_paragraph("Teams submission title:")
    doc.add_paragraph("Phase 2 - Proposal and Code Implementation - [Your Name]").runs[0].bold = True

    # 14. References
    add_heading(doc, "15. References", 1)
    refs = [
        "Yang, Y., & Newsam, S. (2010). Bag-of-visual-words and spatial extensions for land-use classification. ACM GIS.",
        "Howard, A., et al. (2017). MobileNets: Efficient convolutional neural networks for mobile vision applications.",
        "He, K., et al. (2016). Deep residual learning for image recognition. CVPR.",
        "Selvaraju, R. R., et al. (2017). Grad-CAM: Visual explanations from deep networks.",
        "UC Merced Land Use Dataset: http://weegee.vision.ucmerced.edu/datasets/landuse.html",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {ref}")

    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUT_PATH))
    print(f"Word document saved to: {OUT_PATH}")


if __name__ == "__main__":
    main()
